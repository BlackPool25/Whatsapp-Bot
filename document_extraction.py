"""
Document Text Extraction Service
Fast and accurate text extraction from PDFs, DOCX, TXT, CSV and other document formats
"""
import io
from typing import Tuple, Optional, Dict
import mimetypes


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF (fastest method)
    
    Args:
        file_content: PDF file bytes
    
    Returns:
        str: Extracted text content
    """
    try:
        import pymupdf  # PyMuPDF - fastest PDF text extraction
        
        # Open PDF from bytes
        doc = pymupdf.open(stream=file_content, filetype="pdf")
        
        # Extract text from all pages (optimized with sort=True for better layout)
        text_parts = []
        for page_num, page in enumerate(doc):
            # Get text with sorting for better layout preservation
            page_text = page.get_text(sort=True)
            if page_text.strip():
                text_parts.append(page_text)
        
        doc.close()
        
        # Join all pages with double newline separator
        full_text = "\n\n".join(text_parts)
        return full_text.strip()
    
    except ImportError:
        raise ImportError("PyMuPDF (pymupdf) is required for PDF extraction. Install with: pip install pymupdf")
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX using python-docx
    
    Args:
        file_content: DOCX file bytes
    
    Returns:
        str: Extracted text content
    """
    try:
        from docx import Document
        
        # Open DOCX from bytes
        doc = Document(io.BytesIO(file_content))
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts).strip()
    
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from plain text files
    
    Args:
        file_content: Text file bytes
    
    Returns:
        str: Decoded text content
    """
    try:
        # Try UTF-8 first (most common)
        try:
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            # Fallback to latin-1 (handles most edge cases)
            return file_content.decode('latin-1').strip()
    except Exception as e:
        raise Exception(f"Failed to decode text file: {str(e)}")


def extract_text_from_csv(file_content: bytes) -> str:
    """
    Extract text from CSV files
    
    Args:
        file_content: CSV file bytes
    
    Returns:
        str: Formatted CSV content as text
    """
    try:
        import csv
        
        # Decode bytes to string
        content_str = file_content.decode('utf-8')
        
        # Parse CSV
        csv_reader = csv.reader(io.StringIO(content_str))
        
        # Convert to formatted text
        text_parts = []
        for row in csv_reader:
            if any(cell.strip() for cell in row):  # Skip empty rows
                text_parts.append(" | ".join(cell.strip() for cell in row))
        
        return "\n".join(text_parts).strip()
    
    except Exception as e:
        raise Exception(f"Failed to extract text from CSV: {str(e)}")


def extract_text_from_document(
    file_content: bytes, 
    filename: str = None, 
    mime_type: str = None
) -> Tuple[str, Dict[str, any]]:
    """
    Universal document text extraction
    
    Supports: PDF, DOCX, DOC, TXT, CSV, and other text-based formats
    
    Args:
        file_content: File bytes
        filename: Original filename (optional, used for extension detection)
        mime_type: MIME type (optional)
    
    Returns:
        Tuple[str, Dict]: (extracted_text, metadata)
            - extracted_text: Full text content
            - metadata: Dict with extraction info (pages, chars, method, etc.)
    """
    # Determine file type
    extension = None
    if filename and '.' in filename:
        extension = filename.rsplit('.', 1)[1].lower()
    
    if not mime_type and filename:
        mime_type = mimetypes.guess_type(filename)[0]
    
    # Initialize metadata
    metadata = {
        "extraction_method": None,
        "file_type": extension or "unknown",
        "mime_type": mime_type or "unknown",
        "char_count": 0,
        "word_count": 0,
        "success": False,
        "error": None
    }
    
    try:
        extracted_text = ""
        
        # PDF extraction
        if extension == 'pdf' or (mime_type and 'pdf' in mime_type):
            extracted_text = extract_text_from_pdf(file_content)
            metadata["extraction_method"] = "pymupdf"
        
        # DOCX extraction
        elif extension in ['docx', 'doc'] or (mime_type and 'wordprocessingml' in mime_type):
            extracted_text = extract_text_from_docx(file_content)
            metadata["extraction_method"] = "python-docx"
        
        # CSV extraction
        elif extension == 'csv' or (mime_type and 'csv' in mime_type):
            extracted_text = extract_text_from_csv(file_content)
            metadata["extraction_method"] = "csv-parser"
        
        # Plain text extraction (TXT and fallback)
        else:
            extracted_text = extract_text_from_txt(file_content)
            metadata["extraction_method"] = "text-decode"
        
        # Update metadata
        metadata["char_count"] = len(extracted_text)
        metadata["word_count"] = len(extracted_text.split())
        metadata["success"] = True
        
        return extracted_text, metadata
    
    except Exception as e:
        metadata["error"] = str(e)
        metadata["success"] = False
        raise Exception(f"Document extraction failed: {str(e)}")


def is_valid_document_for_detection(text: str, min_chars: int = 20, max_chars: int = 100000) -> Tuple[bool, str]:
    """
    Validate extracted text is suitable for AI detection
    
    Args:
        text: Extracted text
        min_chars: Minimum character count
        max_chars: Maximum character count (to prevent abuse)
    
    Returns:
        Tuple[bool, str]: (is_valid, error_message)
    """
    if not text or not text.strip():
        return False, "No text content found in document"
    
    char_count = len(text.strip())
    
    if char_count < min_chars:
        return False, f"Document too short ({char_count} chars). Minimum {min_chars} characters required."
    
    if char_count > max_chars:
        return False, f"Document too long ({char_count} chars). Maximum {max_chars} characters allowed."
    
    return True, ""


# Example usage and testing
if __name__ == "__main__":
    # Test with sample PDF bytes
    print("Document extraction service initialized")
    print("Supported formats: PDF, DOCX, DOC, TXT, CSV")
    print("Methods: PyMuPDF (PDF), python-docx (DOCX), native (TXT/CSV)")
